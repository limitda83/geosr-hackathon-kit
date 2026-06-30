"""
보고서 생성 서브에이전트 (report_agent.py)

역할:
  - Heat Analysis 시각화 자료(PNG) + 시계열 데이터 + 뉴스 집계를 묶어
    Claude AI가 분석 내러티브를 작성하고 Word/PDF로 저장
  - 사용자 커스텀 옵션으로 포함/제외 섹션, 제목, 기간, 임계값 조절 가능

사용법:
  from agents.report_agent import run

  # 기본 실행
  path = run()

  # 커스텀 실행
  path = run(
      title="2025년 8월 고수온 특보 분석",
      threshold=27.5,
      include_sections=["overview", "timeseries", "hotmap", "news"],
      top_n_regions=10,
      fmt="docx",          # "docx" | "pdf" | "both"
      output_dir="data/results/summary",
  )
"""

from __future__ import annotations

import io
import os
from datetime import datetime
from pathlib import Path
from typing import Literal

import pandas as pd

# ── 경로 상수 ─────────────────────────────────────────────
_ROOT        = Path(__file__).parent.parent
DISASTER_DB  = _ROOT / "data/processed/disaster_db/disaster_events.csv"
REGION_FREQ  = _ROOT / "data/results/frequency/region_frequency.csv"
SST_DIR      = _ROOT / "data"
SST_HOT_IMG  = _ROOT / "data/results/sst_analysis/sst_over28/img"
SST_PERS_DIR = _ROOT / "data/results/sst_analysis/persistence"
SST_TS_DIR   = _ROOT / "data/results/sst_analysis/timeseries"

# 포함 가능한 섹션 목록 (순서 = 보고서 순서, 번들 구조 반영)
ALL_SECTIONS = ["overview", "news", "consec_map", "freq_map", "timeseries", "sst_stats", "hotmap", "conclusion"]
SST_HOTLOW_IMG  = _ROOT / "data/results/hotlowsal/img"
REGION_MAP_PNG  = _ROOT / "data/results/frequency/region_map.png"
REGION_MAP_DIR  = _ROOT / "data/results/frequency"


# ── 데이터 로더 ──────────────────────────────────────────

def _load_news() -> pd.DataFrame:
    return pd.read_csv(DISASTER_DB, encoding="utf-8", parse_dates=["date"])


def _make_region_map_png(freq_df: pd.DataFrame) -> Path:
    """관심지역 지도를 PNG로 저장하고 경로 반환."""
    import plotly.express as px
    df = freq_df.dropna(subset=["lat", "lon"]).copy()
    fig = px.scatter_geo(
        df, lat="lat", lon="lon", size="count",
        hover_name="location", text="location",
        color="count", color_continuous_scale="Teal",
        size_max=30, scope="asia",
        center=dict(lat=36.5, lon=127.5),
    )
    fig.update_geos(
        fitbounds="locations", visible=True, resolution=50,
        showcountries=True, countrycolor="#aaaaaa",
        showland=True, landcolor="#f0f4f8",
        showocean=True, oceancolor="#ddeeff",
    )
    fig.update_layout(
        margin=dict(l=0, r=0, t=0, b=0),
        coloraxis_showscale=False, showlegend=False,
        height=500, width=700,
    )
    fig.update_traces(textposition="top center", textfont_size=9)
    out = REGION_MAP_DIR / "region_map.png"
    out.parent.mkdir(parents=True, exist_ok=True)
    fig.write_image(str(out), scale=2)
    return out


def _load_sst_frames(threshold: float) -> tuple[dict, pd.DataFrame]:
    """지역별 SST CSV 로드 + 통계 DataFrame 반환."""
    frames, stats = {}, []
    for f in sorted(SST_DIR.glob("*_2025*.csv")):
        region = f.stem.split("_")[0]
        if region in ("geocode", "regions", "Tongyeong"):
            continue
        df = pd.read_csv(f, encoding="utf-8", parse_dates=["date"])
        if "sst" not in df.columns:
            continue
        if df.get("source", pd.Series([""])).iloc[0] != "KHOA_OPeNDAP":
            continue
        frames[region] = df
        hot = df["sst"] >= threshold
        cur = mx = 0
        for v in hot:
            cur = cur + 1 if v else 0
            mx = max(mx, cur)
        stats.append({
            "지역": region,
            "평균SST(°C)": round(df["sst"].mean(), 2),
            "최고SST(°C)": round(df["sst"].max(), 2),
            f"고수온일수(≥{threshold}°C)": int(hot.sum()),
            "최장연속(일)": mx,
        })
    stat_df = pd.DataFrame(stats).sort_values(
        f"고수온일수(≥{threshold}°C)", ascending=False
    ) if stats else pd.DataFrame()
    return frames, stat_df


def _load_timeseries() -> pd.DataFrame | None:
    csvs = sorted(SST_TS_DIR.glob("SST_daily_mean_*.csv")) if SST_TS_DIR.exists() else []
    if not csvs:
        return None
    return pd.read_csv(csvs[0], encoding="utf-8-sig", parse_dates=["date"])


def _ai_narrative(context: str, title: str, sections: list[str]) -> str:
    """Claude Haiku로 분석 내러티브 생성. API 키 없으면 기본 텍스트 반환."""
    try:
        import anthropic
        api_key = os.environ.get("ANTHROPIC_API_KEY")
        if not api_key:
            raise ValueError("ANTHROPIC_API_KEY 없음")
        client = anthropic.Anthropic(api_key=api_key)
        section_guide = {
            "overview":   "개요: 분석 목적과 기간 요약",
            "timeseries": "시계열 분석: 일평균 SST 추이와 고수온 발생 패턴",
            "hotmap":     "공간 분포: 고수온(28°C↑) 격자 분포 해석",
            "freq_map":   "누적 빈도: 2일 이상 고수온 누적 지역 특징",
            "consec_map": "연속 지속: 3일 이상 연속 고수온 지역 특징",
            "sst_stats":  "지역별 SST 통계: 피해 위험 지역 분석",
            "news":       "뉴스 분석: 재난 기사 빈도로 본 피해 현황",
        }
        guide = "\n".join(f"- {section_guide[s]}" for s in sections if s in section_guide)
        resp = client.messages.create(
            model="claude-haiku-4-5-20251001",
            max_tokens=2048,
            messages=[{"role": "user", "content":
                f"다음 데이터를 바탕으로 '{title}' 보고서의 분석 내러티브를 작성하세요.\n"
                f"포함 섹션:\n{guide}\n\n데이터:\n{context}\n\n"
                "각 섹션별로 2~3문장의 분석 요약을 작성하세요. 전문적이고 간결하게."}],
        )
        return resp.content[0].text
    except Exception as e:
        return f"[AI 내러티브 생성 생략: {e}]\n분석 결과는 아래 표와 이미지를 참조하세요."


# ── Word 생성 ─────────────────────────────────────────────

def _build_docx(
    title: str,
    sections: list[str],
    news_df: pd.DataFrame,
    sst_stat_df: pd.DataFrame,
    ts_df: pd.DataFrame | None,
    narrative: str,
    threshold: float,
    top_n: int,
    freq_df: pd.DataFrame | None = None,
) -> bytes:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    doc.add_heading(title, 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"생성일: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph(f"분석 기간: 2025-06-01 ~ 2025-08-31  |  고수온 기준: {threshold}°C")
    doc.add_paragraph()

    # AI 내러티브
    if narrative:
        doc.add_heading("분석 요약 (AI)", 1)
        for line in narrative.split("\n"):
            if line.strip():
                doc.add_paragraph(line.strip())
        doc.add_paragraph()

    section_funcs = {
        "overview": lambda: _docx_overview(doc, news_df, sst_stat_df),
        "timeseries": lambda: _docx_timeseries(doc, ts_df),
        "hotmap": lambda: _docx_image(doc, sorted(SST_HOT_IMG.glob("*.png")),
                                      "일별 고수온(28°C↑) 분포 지도", max_imgs=3),
        "freq_map": lambda: _docx_image(doc, sorted(SST_PERS_DIR.glob("SST_HOTFREQ*.png")),
                                        "누적 빈도 분포 지도 (2일↑)"),
        "consec_map": lambda: _docx_image(doc, sorted(SST_PERS_DIR.glob("SST_MAXCONSEC*.png")),
                                          "최장 연속 분포 지도 (3일↑)"),
        "sst_stats": lambda: _docx_sst_stats(doc, sst_stat_df),
        "news": lambda: _docx_news(doc, news_df, top_n, freq_df),
    }

    for sec in sections:
        if sec in section_funcs:
            section_funcs[sec]()

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def _docx_overview(doc, news_df, sst_stat_df):
    from docx import Document
    doc.add_heading("1. 개요", 1)
    top = news_df["location"].value_counts().index[0] if not news_df.empty else "-"
    doc.add_paragraph(
        f"총 {len(news_df):,}건의 재난 뉴스 기사 분석 결과, 가장 많이 언급된 지역은 '{top}'입니다. "
        f"KHOA OPeNDAP 실측 SST 데이터를 기반으로 {len(sst_stat_df)}개 연안 지역의 고수온 현황을 분석하였습니다."
    )


def _docx_timeseries(doc, ts_df):
    from docx.shared import Inches
    doc.add_heading("2. 일평균 SST 시계열", 1)
    # 이미지
    imgs = sorted(SST_TS_DIR.glob("SST_daily_mean_*.png")) if SST_TS_DIR.exists() else []
    if imgs:
        doc.add_picture(str(imgs[0]), width=Inches(5.5))
    # 표
    if ts_df is not None:
        doc.add_paragraph("일평균 SST 데이터 (상위 10일):")
        show = ts_df.sort_values("mean_sst", ascending=False).head(10)
        cols = [c for c in show.columns if c != "valid_cells"]
        t = doc.add_table(rows=1, cols=len(cols), style="Table Grid")
        for i, c in enumerate(cols):
            t.rows[0].cells[i].text = c
        for _, r in show[cols].iterrows():
            row = t.add_row().cells
            for i, v in enumerate(r):
                row[i].text = f"{v:.2f}" if isinstance(v, float) else str(v)[:19]
    doc.add_paragraph()


def _docx_image(doc, imgs: list, heading: str, max_imgs: int = 1):
    from docx.shared import Inches
    if not imgs:
        return
    doc.add_heading(heading, 1)
    for img in imgs[:max_imgs]:
        doc.add_picture(str(img), width=Inches(5.5))
    doc.add_paragraph()


def _docx_sst_stats(doc, sst_stat_df):
    if sst_stat_df.empty:
        return
    doc.add_heading("지역별 SST 통계", 1)
    cols = list(sst_stat_df.columns)
    t = doc.add_table(rows=1, cols=len(cols), style="Table Grid")
    for i, c in enumerate(cols):
        t.rows[0].cells[i].text = c
    for _, r in sst_stat_df.iterrows():
        row = t.add_row().cells
        for i, v in enumerate(r):
            row[i].text = str(v)
    doc.add_paragraph()


def _docx_news(doc, news_df, top_n, freq_df: pd.DataFrame | None = None):
    from docx.shared import Inches
    doc.add_heading("고수온 이벤트 분석결과 (뉴스 크롤링)", 1)
    if freq_df is not None and not freq_df.empty:
        doc.add_paragraph(f"설정 기간 지역별 고수온 재난 발생 횟수 (총 {len(news_df):,}건)")
        show = freq_df.sort_values("count", ascending=False).head(10)
        t = doc.add_table(rows=1, cols=3, style="Table Grid")
        for i, h in enumerate(["순위", "지역", "발생 건수"]):
            t.rows[0].cells[i].text = h
        for i, (_, r) in enumerate(show.iterrows(), 1):
            row = t.add_row().cells
            row[0].text = str(i)
            row[1].text = str(r["location"])
            row[2].text = str(r["count"])
        doc.add_paragraph()
        # 관심지역 지도
        try:
            map_png = _make_region_map_png(freq_df)
            doc.add_heading("관심지역 분포 지도", 1)
            doc.add_picture(str(map_png), width=Inches(5.5))
        except Exception:
            pass
        doc.add_paragraph()


# ── PDF 생성 ─────────────────────────────────────────────

def _build_pdf(
    title: str,
    sections: list[str],
    news_df: pd.DataFrame,
    sst_stat_df: pd.DataFrame,
    ts_df: pd.DataFrame | None,
    narrative: str,
    threshold: float,
    top_n: int,
    freq_df: pd.DataFrame | None = None,
) -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.lib import colors
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer,
                                    Table, TableStyle, HRFlowable,
                                    Image as RLImage, PageBreak)
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    font_name = "Helvetica"
    for fp in ["C:/Windows/Fonts/malgun.ttf", "C:/Windows/Fonts/NanumGothic.ttf",
               "/usr/share/fonts/truetype/nanum/NanumGothic.ttf"]:
        if Path(fp).exists():
            try:
                pdfmetrics.registerFont(TTFont("KorFont", fp))
                font_name = "KorFont"
            except Exception:
                pass
            break

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
                            leftMargin=20*mm, rightMargin=20*mm,
                            topMargin=20*mm, bottomMargin=20*mm)

    kor  = ParagraphStyle("kor",  fontName=font_name, fontSize=9,  leading=15)
    h1   = ParagraphStyle("h1",   fontName=font_name, fontSize=15, leading=20, spaceAfter=4,
                           textColor=colors.HexColor("#041529"))
    h2   = ParagraphStyle("h2",   fontName=font_name, fontSize=12, leading=17, spaceAfter=3,
                           textColor=colors.HexColor("#0a3660"))
    bold = ParagraphStyle("bold", fontName=font_name, fontSize=9,  leading=15, fontStyle="bold")

    def _table(data, col_widths=None):
        t = Table(data, colWidths=col_widths)
        t.setStyle(TableStyle([
            ("BACKGROUND", (0,0), (-1,0), colors.HexColor("#0a3660")),
            ("TEXTCOLOR",  (0,0), (-1,0), colors.white),
            ("FONTNAME",   (0,0), (-1,-1), font_name),
            ("FONTSIZE",   (0,0), (-1,-1), 8),
            ("ROWBACKGROUNDS", (0,1), (-1,-1),
             [colors.white, colors.HexColor("#eef3f8")]),
            ("GRID", (0,0), (-1,-1), 0.3, colors.HexColor("#cccccc")),
            ("ALIGN", (1,0), (-1,-1), "CENTER"),
            ("VALIGN", (0,0), (-1,-1), "MIDDLE"),
            ("TOPPADDING",   (0,0), (-1,-1), 3),
            ("BOTTOMPADDING",(0,0), (-1,-1), 3),
        ]))
        return t

    story = []

    # 표지
    story += [
        Spacer(1, 8*mm),
        Paragraph(title, h1),
        Paragraph(f"생성일: {datetime.now().strftime('%Y-%m-%d %H:%M')}  |  "
                  f"기간: 2025-06-01 ~ 2025-08-31  |  고수온 기준: {threshold}°C", kor),
        HRFlowable(width="100%", thickness=1,
                   color=colors.HexColor("#00c2d4"), spaceAfter=6),
    ]

    if narrative:
        story.append(Paragraph("AI 분석 요약", h2))
        for line in narrative.split("\n"):
            if line.strip():
                story.append(Paragraph(line.strip(), kor))
        story.append(Spacer(1, 4*mm))

    # SST 트렌드 수치 계산
    sst_start = sst_end = sst_max = None
    if ts_df is not None:
        _s = ts_df.sort_values("date")
        sst_start = round(_s.iloc[0]["mean_sst"], 1)
        sst_end   = round(_s.iloc[-1]["mean_sst"], 1)
        sst_max   = round(ts_df["mean_sst"].max(), 1)

    top_region = news_df["location"].value_counts().index[0] if not news_df.empty else "-"
    top5 = "·".join(
        (freq_df.sort_values("count", ascending=False)["location"].head(5).tolist()
         if freq_df is not None and not freq_df.empty else [])
    ) or top_region

    sec_order = {s: i for i, s in enumerate(ALL_SECTIONS)}
    for sec in sorted(sections, key=lambda s: sec_order.get(s, 99)):

        if sec == "overview":
            story += [
                Paragraph("개요", h2),
                Paragraph(
                    f"재난 뉴스(고수온) 발생 빈도로 관심지역을 도출하고, "
                    f"국가해양위성센터 해수면온도(SST) 자료로 해당 해역의 "
                    f"고수온 발생·지속·트렌드를 분석함.", kor),
                Paragraph(
                    f"뉴스 이벤트 총 <b>{len(news_df):,}건</b> / "
                    f"해수면온도 2025.7.1~8.31(62일) / 고수온 기준 <b>{threshold}°C</b>.", kor),
                Spacer(1, 4*mm),
            ]

        elif sec == "news":
            if freq_df is not None and not freq_df.empty:
                story.append(Paragraph("고수온 이벤트 분석결과 (뉴스 크롤링)", h2))
                story.append(Paragraph(
                    f"설정 기간 동안 지역별 고수온 재난 발생 횟수(총 {len(news_df):,}건). "
                    f"{top5} 등 남해안과 제주에 발생이 집중됨.", kor))
                # 빈도 표
                show = freq_df.sort_values("count", ascending=False).head(10)
                tdata = [["순위", "지역", "발생 건수"]]
                for i, (_, r) in enumerate(show.iterrows(), 1):
                    tdata.append([str(i), str(r["location"]), str(r["count"])])
                story.append(_table(tdata, [20*mm, 80*mm, 70*mm]))
                story.append(Spacer(1, 3*mm))
                # 관심지역 지도
                try:
                    map_png = _make_region_map_png(freq_df)
                    story.append(Paragraph("관심지역 분포 지도", h2))
                    story.append(RLImage(str(map_png), width=160*mm, height=110*mm))
                except Exception:
                    pass
                story.append(Spacer(1, 4*mm))

        elif sec == "consec_map":
            imgs = sorted(SST_PERS_DIR.glob("SST_MAXCONSEC*.png"))
            story.append(Paragraph("고수온 지속일수 분석 (2025.7.1~8.31)", h2))
            story.append(Paragraph(
                "남해안 연안을 중심으로 28℃ 이상이 여러 날 연속 지속된 해역이 형성됨. "
                "연속 지속일수가 긴 격자일수록 양식생물 피해 위험이 누적되어 우선 관리 대상.", kor))
            if imgs:
                story.append(RLImage(str(imgs[0]), width=160*mm, height=80*mm))
            story.append(Spacer(1, 4*mm))

        elif sec == "freq_map":
            imgs = sorted(SST_PERS_DIR.glob("SST_HOTFREQ*.png"))
            story.append(Paragraph("고수온 빈도 공간 분석 (2025.7.1~8.31)", h2))
            story.append(Paragraph(
                f"남해안·제주 연안에서 고수온 발생 빈도가 높게 누적되며, "
                f"뉴스 기반 관심지역({top5})과 공간적으로 일치함.", kor))
            if imgs:
                story.append(RLImage(str(imgs[0]), width=160*mm, height=80*mm))
            story.append(Spacer(1, 4*mm))

        elif sec == "timeseries":
            imgs = sorted(SST_TS_DIR.glob("SST_daily_mean_*.png")) if SST_TS_DIR.exists() else []
            story.append(Paragraph("평균 해수면온도 트렌드 (2025.7.1~8.31)", h2))
            if sst_start is not None:
                diff = round(sst_end - sst_start, 1)
                story.append(Paragraph(
                    f"7월 초 평균 {sst_start}℃에서 8월 말 {sst_end}℃로 약 {diff}℃ 지속 상승, "
                    f"격자 평균 최고 {sst_max}℃. 여름철 해수면온도의 뚜렷한 상승 추세.", kor))
            if imgs:
                story.append(RLImage(str(imgs[0]), width=160*mm, height=55*mm))
            story.append(Spacer(1, 4*mm))

        elif sec == "sst_stats" and not sst_stat_df.empty:
            story.append(Paragraph("지역별 SST 통계", h2))
            cols_s = list(sst_stat_df.columns)
            tdata  = [cols_s] + [[str(v) for v in row] for row in sst_stat_df.values]
            story.append(_table(tdata, [170*mm / len(cols_s)] * len(cols_s)))
            story.append(Spacer(1, 4*mm))

        elif sec == "hotmap":
            imgs = sorted(SST_HOT_IMG.glob("*.png"))[-3:]
            if imgs:
                story.append(Paragraph("일별 고수온 분포 지도 (최근 3일)", h2))
                for img in imgs:
                    date_str = img.stem.split("_U")[-1].replace("_HOT28", "")
                    date_fmt = f"{date_str[:4]}-{date_str[4:6]}-{date_str[6:]}"
                    story.append(Paragraph(date_fmt, kor))
                    story.append(RLImage(str(img), width=160*mm, height=70*mm))
                story.append(Spacer(1, 4*mm))

        elif sec == "conclusion":
            story += [
                Paragraph("결론 및 시사점", h2),
                Paragraph(
                    f"뉴스 발생 빈도와 위성 고수온 분석이 남해안·제주에서 공간적으로 일치 "
                    f"→ 관심지역({top5}) 자동 도출의 타당성을 확인함.", kor),
                Paragraph(
                    "고수온 장기 지속 해역과 저염분 중첩 해역을 우선 경계 대상으로 설정하여 "
                    "양식어가 피해 예방에 선제적으로 대응 가능.", kor),
                Spacer(1, 3*mm),
                Paragraph("※ 출처: 한국언론진흥재단·빅카인즈(뉴스 메타데이터), "
                           "국가해양위성센터 KHOA L4 해수면온도", kor),
                Spacer(1, 4*mm),
            ]

    doc.build(story)
    buf.seek(0)
    return buf.read()


# ── 공개 인터페이스 ───────────────────────────────────────

def run(
    title: str = "고수온 연안재해 분석 보고서",
    threshold: float = 28.0,
    include_sections: list[str] | None = None,
    top_n_regions: int = 20,
    fmt: Literal["docx", "pdf", "both"] = "both",
    output_dir: str = "data/results/summary",
    use_ai: bool = True,
) -> dict:
    """
    보고서를 생성하고 bytes와 저장 경로를 함께 반환한다.

    Parameters
    ----------
    title            : 보고서 제목
    threshold        : 고수온 기준 온도 (기본 28.0°C)
    include_sections : 포함할 섹션 목록 (None이면 전체)
                       선택 가능: overview, timeseries, hotmap, freq_map,
                                  consec_map, sst_stats, news
    top_n_regions    : 뉴스 섹션에 포함할 최근 기사 수
    fmt              : "docx" | "pdf" | "both"
    output_dir       : 저장 폴더 (None이면 저장 안 함)
    use_ai           : Claude AI 내러티브 생성 여부

    Returns
    -------
    dict: {
        "docx": {"bytes": bytes, "path": str} | None,
        "pdf":  {"bytes": bytes, "path": str} | None,
        "filename_stem": str,
    }
    """
    sections = include_sections if include_sections else ALL_SECTIONS[:]

    stamp = datetime.now().strftime("%Y%m%d_%H%M")

    print("[report_agent] 데이터 로딩...")
    news_df        = _load_news()
    _, sst_stat_df = _load_sst_frames(threshold)
    ts_df          = _load_timeseries()
    freq_df        = pd.read_csv(REGION_FREQ, encoding="utf-8") if REGION_FREQ.exists() else None

    narrative = ""
    if use_ai:
        print("[report_agent] AI 내러티브 생성 중...")
        ctx_lines = [f"뉴스 기사 수: {len(news_df)}건"]
        if not sst_stat_df.empty:
            ctx_lines.append(sst_stat_df.to_string(index=False))
        if ts_df is not None:
            ctx_lines.append(
                f"일평균 SST 범위: {ts_df['mean_sst'].min():.1f}~{ts_df['mean_sst'].max():.1f}°C"
            )
        narrative = _ai_narrative("\n".join(ctx_lines), title, sections)

    result: dict = {"docx": None, "pdf": None, "filename_stem": f"report_{stamp}"}

    if fmt in ("docx", "both"):
        print("[report_agent] Word 생성 중...")
        docx_bytes = _build_docx(title, sections, news_df, sst_stat_df,
                                 ts_df, narrative, threshold, top_n_regions, freq_df)
        entry: dict = {"bytes": docx_bytes, "path": None}
        if output_dir:
            out_dir = Path(output_dir)
            out_dir.mkdir(parents=True, exist_ok=True)
            p = out_dir / f"report_{stamp}.docx"
            p.write_bytes(docx_bytes)
            entry["path"] = str(p)
            print(f"  → {p}")
        result["docx"] = entry

    if fmt in ("pdf", "both"):
        print("[report_agent] PDF 생성 중...")
        pdf_bytes = _build_pdf(title, sections, news_df, sst_stat_df,
                               ts_df, narrative, threshold, top_n_regions, freq_df)
        entry = {"bytes": pdf_bytes, "path": None}
        if output_dir:
            out_dir = Path(output_dir)
            out_dir.mkdir(parents=True, exist_ok=True)
            p = out_dir / f"report_{stamp}.pdf"
            p.write_bytes(pdf_bytes)
            entry["path"] = str(p)
            print(f"  → {p}")
        result["pdf"] = entry

    print(f"[report_agent] 완료")
    return result


if __name__ == "__main__":
    # 직접 실행 예시
    result = run(
        title="2025년 하절기 고수온 연안재해 분석 보고서",
        threshold=28.0,
        include_sections=["overview", "timeseries", "hotmap", "freq_map", "consec_map", "sst_stats", "news"],
        top_n_regions=20,
        fmt="both",
    )
    print(result)
